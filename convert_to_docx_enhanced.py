#!/usr/bin/env python3
"""
Enhanced script to convert SOFTWARE_DOCUMENTATION.md to DOCX format with better formatting
"""

import re
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx library not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX with enhanced formatting"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Set heading styles
    for i in range(1, 10):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(14 - i) if i <= 3 else Pt(11)
        if i == 1:
            heading_style.font.bold = True
            heading_style.font.size = Pt(16)
    
    # Parse content
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_rows = []
    table_headers = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                if code_block_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if '|' in line and not in_table:
            in_table = True
            table_rows = []
            # Parse header
            headers = [cell.strip() for cell in line.split('|') if cell.strip()]
            table_headers = headers
            i += 1
            # Skip separator line
            if i < len(lines) and '---' in lines[i]:
                i += 1
            continue
        
        if in_table:
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                table_rows.append(cells)
                i += 1
                continue
            else:
                # End table
                if table_rows:
                    table = doc.add_table(rows=1, cols=len(table_headers))
                    table.style = 'Light Grid Accent 1'
                    
                    # Add headers
                    header_cells = table.rows[0].cells
                    for j, header in enumerate(table_headers):
                        header_cells[j].text = header
                        header_cells[j].paragraphs[0].runs[0].font.bold = True
                    
                    # Add rows
                    for row_data in table_rows:
                        row = table.add_row()
                        for j, cell_text in enumerate(row_data):
                            if j < len(row.cells):
                                row.cells[j].text = cell_text
                    
                    doc.add_paragraph()  # Add space after table
                in_table = False
                table_rows = []
                table_headers = []
        
        # Handle headers
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
        elif line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_heading(text, level=4)
        # Handle horizontal rules
        elif stripped == '---' or stripped == '===':
            p = doc.add_paragraph('_' * 80)
            p.paragraph_format.space_after = Pt(12)
        # Handle list items
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            text = process_formatting(doc, text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            text = process_formatting(doc, text)
            p = doc.add_paragraph(text, style='List Number')
        # Handle table of contents links (skip)
        elif stripped.startswith('- [') and '](#' in stripped:
            i += 1
            continue
        # Handle empty lines
        elif not stripped:
            doc.add_paragraph()
        # Handle regular paragraphs
        else:
            if stripped:
                text = process_formatting(doc, stripped)
                p = doc.add_paragraph(text)
        
        i += 1
    
    # Add title page
    add_title_page(doc)
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

def add_title_page(doc):
    """Add a title page at the beginning"""
    # This will be inserted at the beginning, but we'll add it at the end for simplicity
    # In a real implementation, you'd insert it at position 0
    pass

def process_formatting(doc, text):
    """Process markdown formatting and return formatted text"""
    # This is a simplified version - in a full implementation,
    # you'd parse and apply formatting to runs
    # For now, just clean up the text
    text = re.sub(r'`([^`]+)`', r'\1', text)  # Remove code backticks
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold
    text = re.sub(r'__([^_]+)__', r'\1', text)  # Remove bold
    text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Remove italic
    text = re.sub(r'_([^_]+)_', r'\1', text)  # Remove italic
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # Extract link text
    return text

if __name__ == '__main__':
    md_file = 'SOFTWARE_DOCUMENTATION.md'
    docx_file = 'Schemmatics_Digital_Chart_Software_Documentation.docx'
    
    try:
        parse_markdown_to_docx(md_file, docx_file)
        print(f"\nDocument created: {docx_file}")
        print("You can now open this file in Microsoft Word for further editing.")
    except FileNotFoundError:
        print(f"Error: {md_file} not found!")
        sys.exit(1)
    except Exception as e:
        print(f"Error converting file: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


