#!/usr/bin/env python3
"""
Script to convert SOFTWARE_DOCUMENTATION.md to DOCX format
"""

import re
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx library not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set hyperlink style
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse content
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_block_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_block_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    p.style = 'No Spacing'
                code_block_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=4)
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 80)
        # Handle list items
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            # Check for bold/italic
            text = process_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = process_formatting(text)
            p = doc.add_paragraph(text, style='List Number')
        # Handle table of contents links (skip, will be auto-generated)
        elif line.strip().startswith('- [') and '](#' in line:
            continue
        # Handle empty lines
        elif not line.strip():
            doc.add_paragraph()
        # Handle regular paragraphs
        else:
            if line.strip():
                text = process_formatting(line.strip())
                p = doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

def process_formatting(text):
    """Process markdown formatting (bold, italic, code, links)"""
    # Handle inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # Handle bold
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)
    
    # Handle italic
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    # Handle links (extract text only for now)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    return text

if __name__ == '__main__':
    md_file = 'SOFTWARE_DOCUMENTATION.md'
    docx_file = 'Schemmatics_Digital_Chart_Software_Documentation.docx'
    
    try:
        parse_markdown_to_docx(md_file, docx_file)
    except FileNotFoundError:
        print(f"Error: {md_file} not found!")
        sys.exit(1)
    except Exception as e:
        print(f"Error converting file: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


